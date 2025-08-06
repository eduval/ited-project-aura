import os
from docx import Document
from openpyxl import load_workbook
import datetime

def generate_transcripts(input_excel_file, word_output_dir, template_docx, students):
    current_date = datetime.datetime.now().strftime("%Y-%m-%d")
    os.makedirs(word_output_dir, exist_ok=True)
    print("Student info received:", students)
    wb = load_workbook(input_excel_file, data_only=True)

     # Construir el mapa de estudiantes usando la lista students
    student_info_map = {}
    for student in students:
        print("Processing student:", student)
        student_no = str(student.get("Student No")).strip()
        if student_no:
            student_info_map[student_no] = student
    print("Student info map created with", student_info_map)     
    # Process all other sheets as student transcripts
    for sheet_name in wb.sheetnames:
        print("Processing sheet:", sheet_name)
        ws = wb[sheet_name]
        student_no = sheet_name

        if student_no not in student_info_map:
            print(f"Missing student metadata for {student_no}, skipping...")
            continue

        data = student_info_map[student_no]
        print("Student data found:", data)
        # Extract metadata
        full_name = f"{data.get('First Name', '')} {data.get('Last Name', '')}".strip()
        program = str(data.get("Program", "Unknown"))
        start_date = str(data.get("Program start date", "Unknown"))
        dob = str(data.get("DOB", "Unknown"))
        end_date = str(data.get("Program End Date", "N/A"))

        # Load template and replace placeholders
        doc = Document(os.path.join(os.path.dirname(__file__), "template.docx"))
        label_map = {
            "Program:": program,
            "Student Name:": full_name,
            "Student Number:": student_no,
            "DOB:": dob,
            "Program Start Date:": start_date,
            "Program End Date:": end_date,
            # "Date: ": current_date
        }

        for para in doc.paragraphs:
            for label, value in label_map.items():
                if label in para.text:
                    for run in para.runs:
                        if label in run.text:
                            run.text = f"{label} {value}"

        # Find paragraph after which to insert the table
        insert_after = None
        address_str = "Address: 2560 E 48th Ave, Vancouver, BC V5S 1G4"
        for idx, para in enumerate(doc.paragraphs):
            if address_str in para.text:
                insert_after = idx
                break

        # Build transcript table
        table = doc.add_table(rows=1, cols=6)
        table.style = "Style2"
        headers = ["Course Code","Course Name","Course Start Date","Course End Date","Attendance  (%)","Grade  (%)"]
        for i in range(6):
            table.rows[0].cells[i].text = headers[i]

        for row in ws.iter_rows(min_row=2, values_only=True):
            if not row[0]:
                continue
            cells = table.add_row().cells
            for j in range(min(6, len(row))):
                cells[j].text = str(row[j]) if row[j] is not None else ""

        # Insert table after address paragraph or at end
        if insert_after is not None:
            doc._body._element.remove(table._element)
            doc.paragraphs[insert_after]._element.addnext(table._element)
        else:
            print(f"Address line not found in Word template for {student_no}")
            doc.add_paragraph("Transcript Table (position unknown):")
            doc._body._element.append(table._element)

        # Save output file
        clean_student_no = student_no.replace("CT", "")
        output_path = os.path.join(word_output_dir, f"iTED-2025-{clean_student_no}.docx")
        doc.save(output_path)
        print(f"Saved: {output_path}")

    print("All transcripts generated :)!")
